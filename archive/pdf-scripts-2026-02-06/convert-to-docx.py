from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Novara AI Opportunities', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Executive Decision Brief')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

meta = doc.add_paragraph('February 2026 | Confidential')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.italic = True

doc.add_paragraph()

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph('This document presents 10 AI opportunities for Novara, each summarized at a level sufficient for executive go/no-go decisions. Opportunities are organized by strategic category and include complexity estimates, dependencies, competitive context, and recommendations.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Complexity: ').bold = True
p.add_run('Low (weeks) | Medium (1-2 quarters) | High (2+ quarters)')
p = doc.add_paragraph()
p.add_run('Impact: ').bold = True
p.add_run('Low | Medium | High')
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run('🟢 Proceed | 🟡 Evaluate | 🔴 Defer')

# Competitive Landscape
doc.add_heading('Competitive Landscape Overview', level=1)

doc.add_heading('Major EHS Software Players', level=2)
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Company'
hdr[1].text = 'AI Maturity'
hdr[2].text = 'Key AI Features'
hdr[3].text = 'Gaps'

data = [
    ('Intelex (Fortive)', 'Medium', 'Document search, incident trend analytics, NLP for reports', 'No real-time monitoring, weak hardware integration'),
    ('Enablon (Wolters Kluwer)', 'Medium', 'Risk prediction, compliance mapping, incident classification', 'Dated UX, limited real-time capabilities'),
    ('VelocityEHS', 'Medium-High', '"Vera" AI assistant, ergonomics analysis, chemical risk AI', 'Less mature on environmental/compliance depth'),
    ('SafetyCulture', 'Medium', 'Smart inspection templates, issue prediction, photo annotation', 'Enterprise complexity gaps, regulatory depth'),
    ('Benchmark Gensuite', 'Low-Medium', 'Compliance intelligence, leading indicators, auto-reporting', 'Dated UX, shallow AI integration'),
    ('EHS Insight', 'Medium', '"EHS Insight Copilot" AI, incident management, compliance tracking', 'Newer entrant; enterprise depth uncertain'),
    ('HammerTech', 'Low-Medium', 'Construction safety focus, permit management, safety observation tools', 'Narrow industry focus; less AI maturity'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_heading('Computer Vision / PPE Detection Specialists', level=2)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Company'
hdr[1].text = 'Focus'
hdr[2].text = 'Funding'
hdr[3].text = 'Threat Level'

data = [
    ('Intenseye', 'Real-time safety analytics from CCTV', '$64M raised', 'High — strong enterprise sales, expanding scope'),
    ('Protex AI', 'PPE and safety zone monitoring', 'Private', 'Medium — purpose-built but narrow'),
    ('Voxel', 'Cloud video analytics platform', 'Series B', 'Medium — platform approach, less EHS-specific'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_heading('Wearables Players', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Company'
hdr[1].text = 'Product'
hdr[2].text = 'Status'

data = [
    ('Kinetic', 'Ergonomic risk wearable', 'Proven ROI data, warehouse/logistics focus'),
    ('Soter Analytics', 'SoterCoach + analytics', 'Acquiring competitors (StrongArm)'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_heading('Market Dynamics', level=2)
doc.add_paragraph('• EHS Software Market: ~$8-10B globally, 10-12% CAGR')
doc.add_paragraph('• AI in EHS Penetration: ~15-20% of tools have meaningful AI (growing fast)')
doc.add_paragraph('• Key Trend: Everyone is adding "AI-powered" features — differentiation window closing')
doc.add_paragraph('• Consolidation: PE roll-ups active; Voxel is a potential acquisition target')

# Competitor Outlook
doc.add_heading('What We Expect Competitors To Do (12-Month Outlook)', level=1)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Competitor'
hdr[1].text = 'Expected Moves'
hdr[2].text = 'Threat to Novara'

data = [
    ('Intelex/Enablon', 'Bolt-on LLM chatbots, generic AI wrappers', 'Low — will be shallow integrations'),
    ('VelocityEHS', 'Expand "Vera" assistant capabilities, add more CV features', 'Medium — already have AI brand recognition'),
    ('SafetyCulture', 'Photo hazard detection in iAuditor, mobile-first AI', 'High — fast iteration, SMB momentum'),
    ('Intenseye', 'Expand beyond PPE to behavior analysis, near-miss detection', 'High — could move up the stack'),
    ('Amazon/AWS', 'May productize internal warehouse safety CV as AWS service', 'Medium — would commoditize CV capabilities'),
    ('Microsoft/Google', 'Copilot/Duet EHS templates, "good enough" AI for small orgs', 'Low — won\'t be deep, but raises baseline expectations'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

# Industry Trends
doc.add_paragraph()
doc.add_heading('Industry Trends (12-Month Horizon)', level=1)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Trend'
hdr[1].text = 'What It Means'
hdr[2].text = 'Our Play'

data = [
    ('LLM Commoditization', 'AI features become table stakes, not differentiators', 'Ship fast while it\'s still special; depth beats breadth'),
    ('Edge AI Maturity', 'Real-time CV gets cheaper, easier to deploy', 'Hazard detection on standard cameras becomes viable'),
    ('Regulatory AI Scrutiny', 'EU AI Act, OSHA watching AI use in safety-critical contexts', 'Build with explainability/auditability from day 1'),
    ('ESG Integration', 'Safety metrics feeding into sustainability/investor reports', 'Connect safety data to ESG dashboards'),
    ('Wearables Normalization', 'Worker pushback decreasing, unions more accepting', 'Hardware partnerships (Kinetic, Soter) become valuable'),
    ('Consolidation', 'PE-backed players acquiring point solutions', 'Build capabilities that make us an acquirer, not a target'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

# The 10 Opportunities
doc.add_page_break()
doc.add_heading('The 10 Opportunities', level=1)

opportunities = [
    {
        'num': 1,
        'title': 'Policy Interpretation Engine',
        'tldr': 'AI that answers compliance questions using your actual policies. Saves 15-30 min per query. Build after Document Ingestion.',
        'what': 'Natural language Q&A system that lets users ask questions about compliance documents and get accurate, cited answers.',
        'example': '"Does OSHA 1910.147 lockout/tagout apply to our California facility?" → Returns relevant policy sections with citations.',
        'tech': 'RAG + LLM + document chunking + citation tracking',
        'differentiation': 'Deep RAG over customer-specific documents with citation links. "It actually knows our policies" vs competitors\' generic Q&A.',
        'complexity': 'Medium',
        'impact': 'High',
        'rec': '🟢 Proceed'
    },
    {
        'num': 2,
        'title': 'Document Ingestion & Normalization',
        'tldr': 'Turn uploaded PDFs and docs into searchable, structured data. Foundation for everything else. Start here.',
        'what': 'Automated pipeline that extracts structured data from uploaded documents (PDFs, Word, scanned images) and normalizes them into a searchable, tagged format.',
        'example': 'Upload a 50-page safety manual → System extracts sections, tags by topic, identifies key entities (regulations, dates, roles).',
        'tech': 'Document AI (layout detection) + entity extraction + classification + OCR',
        'differentiation': 'First mover on EHS-specific document AI. Extract meaning, not just text.',
        'complexity': 'Medium',
        'impact': 'High',
        'rec': '🟢 Proceed (First)'
    },
    {
        'num': 3,
        'title': 'Incident Summarization',
        'tldr': 'One-click executive summaries from incident reports. Low complexity, high visibility. Quick win.',
        'what': 'Automatically generates executive summaries from detailed incident reports, extracting key facts, root causes, and action items.',
        'example': '3-page incident report → 3-paragraph executive summary with: what happened, who was involved, what actions were taken.',
        'tech': 'LLM summarization + template extraction + key entity highlighting',
        'differentiation': 'Low risk, high visibility. "One-click executive brief" is compelling demo.',
        'complexity': 'Low',
        'impact': 'Medium',
        'rec': '🟢 Proceed (Quick Win)'
    },
    {
        'num': 4,
        'title': 'Root Cause Clustering',
        'tldr': 'ML finds hidden patterns across incidents ("17 incidents share inadequate training — all night shift, Plant B"). Catches issues 60-90 days earlier. Major differentiator.',
        'what': 'ML-driven analysis that identifies patterns and clusters across historical incidents to surface systemic issues.',
        'example': '"17 incidents this quarter share \'inadequate training\' as a contributing factor — concentrated in night shift at Plant B"',
        'tech': 'Embedding + clustering + trend analysis + visualization',
        'differentiation': 'Major differentiation. Nobody is doing real ML on incident text to find hidden patterns. This is "predictive safety."',
        'complexity': 'Medium',
        'impact': 'High',
        'rec': '🟢 Proceed'
    },
    {
        'num': 5,
        'title': 'Corrective Action Drafting',
        'tldr': 'AI suggests corrective actions based on what worked for similar past incidents. Cuts drafting time 50%. Build after we have CAPA history.',
        'what': 'AI suggests appropriate corrective and preventive actions (CAPA) based on incident type, severity, and historical precedent.',
        'example': 'Slip-and-fall incident → AI suggests: "1) Install anti-slip mats, 2) Increase cleaning frequency, 3) Add warning signage"',
        'tech': 'Few-shot learning + historical pattern matching + template generation',
        'differentiation': 'Depends on historical CAPA data.',
        'complexity': 'Medium',
        'impact': 'Medium',
        'rec': '🟡 Evaluate'
    },
    {
        'num': 6,
        'title': 'Audit Prep Automation',
        'tldr': 'Auto-gathers evidence, pre-fills responses, flags gaps before audits. "Never scramble for an audit again." High customer demand.',
        'what': 'Automated system that gathers evidence, pre-fills audit responses, and flags gaps before scheduled audits.',
        'example': 'OSHA audit in 2 weeks → System identifies 52 required evidence items, auto-collects 47, flags 5 that need attention.',
        'tech': 'Document retrieval + checklist matching + gap identification',
        'differentiation': 'Audit anxiety is universal. "Never scramble for an audit again" is powerful message.',
        'complexity': 'Medium',
        'impact': 'High',
        'rec': '🟢 Proceed'
    },
    {
        'num': 7,
        'title': 'Customer-Specific Policy Mapping',
        'tldr': 'Connects generic regulations to your specific SOPs. "What\'s our procedure for X?" → Exact document and section. Build after doc foundation is solid.',
        'what': 'Maps generic regulatory requirements to customer-specific procedures and documents.',
        'example': 'User asks about OSHA respiratory protection → System returns: "See your SOP-RES-001, sections 4.2 and 4.5"',
        'tech': 'Entity linking + document cross-referencing + semantic matching',
        'differentiation': 'Depends on Document Ingestion and per-customer doc corpus.',
        'complexity': 'Medium',
        'impact': 'Medium',
        'rec': '🟡 Evaluate'
    },
    {
        'num': 8,
        'title': 'Questionnaire & RFP Automation',
        'tldr': 'AI drafts 80% of compliance questionnaire responses from past answers. Clear ROI but not core EHS. Consider Loopio integration.',
        'what': 'AI drafts responses to compliance questionnaires and RFPs by drawing from the organization\'s knowledge base.',
        'example': 'Customer receives 200-question safety questionnaire → AI drafts 80% of responses from existing policies and past answers.',
        'tech': 'Question classification + knowledge retrieval + response generation',
        'differentiation': 'Clear ROI but not core to EHS differentiation. Consider Loopio/Responsive integration.',
        'complexity': 'Low-Medium',
        'impact': 'Medium',
        'rec': '🟡 Evaluate'
    },
    {
        'num': 9,
        'title': 'Smart Onboarding Assistant',
        'tldr': 'AI-guided setup for new customers. Reduces implementation time 30-50%. Nice to have, not a differentiator.',
        'what': 'Conversational AI that guides new users/customers through system setup with contextual help and auto-configuration.',
        'example': 'New customer → AI asks about industry, size, key regulations → Auto-configures relevant modules, templates, and workflows.',
        'tech': 'Decision tree + LLM conversation + configuration automation',
        'differentiation': 'Reduces implementation burden, accelerates time-to-value. Not a competitive differentiator.',
        'complexity': 'Low',
        'impact': 'Medium',
        'rec': '🟡 Evaluate'
    },
    {
        'num': 10,
        'title': 'Hazard Detection (Computer Vision)',
        'tldr': 'AI analyzes photos/video to spot hazards and PPE violations. High complexity but high impact. Competitive race with Intenseye/SafetyCulture. Already in progress.',
        'what': 'AI-powered analysis of photos and video to automatically identify workplace hazards, PPE compliance issues, and unsafe conditions.',
        'example': 'Worker uploads jobsite photo → AI detects: missing hard hat, unguarded machinery, trip hazard (cable across walkway).',
        'tech': 'Object detection (YOLO/similar) + edge inference + alert system + integration with incident management',
        'differentiation': 'Bridge between EHS software (us) and CV point solutions. Integrated hazard detection within compliance workflow.',
        'complexity': 'High',
        'impact': 'High',
        'rec': '🟢 Proceed'
    },
]

for opp in opportunities:
    doc.add_heading(f"{opp['num']}. {opp['title']}", level=2)
    
    # TLDR
    p = doc.add_paragraph()
    p.add_run('TLDR: ').bold = True
    p.add_run(opp['tldr']).italic = True
    
    # Details
    p = doc.add_paragraph()
    p.add_run('What: ').bold = True
    p.add_run(opp['what'])
    
    p = doc.add_paragraph()
    p.add_run('Example: ').bold = True
    p.add_run(opp['example'])
    
    p = doc.add_paragraph()
    p.add_run('Tech: ').bold = True
    p.add_run(opp['tech'])
    
    p = doc.add_paragraph()
    p.add_run('Our Differentiation: ').bold = True
    p.add_run(opp['differentiation'])
    
    # Badges
    p = doc.add_paragraph()
    p.add_run(f"Complexity: {opp['complexity']} | Impact: {opp['impact']} | {opp['rec']}")
    
    doc.add_paragraph()

# Summary Matrix
doc.add_page_break()
doc.add_heading('Summary Matrix', level=1)

table = doc.add_table(rows=11, cols=7)
table.style = 'Table Grid'
hdr = table.rows[0].cells
headers = ['#', 'Opportunity', 'TLDR', 'Complexity', 'Impact', 'Urgency', 'Rec']
for i, h in enumerate(headers):
    hdr[i].text = h

matrix_data = [
    ('1', 'Policy Interpretation', 'AI answers compliance questions from your docs', 'Medium', 'High', 'Medium', '🟢'),
    ('2', 'Document Ingestion', 'Turn PDFs into searchable, structured data', 'Medium', 'High', 'High', '🟢'),
    ('3', 'Incident Summarization', 'One-click executive summaries', 'Low', 'Medium', 'Low', '🟢'),
    ('4', 'Root Cause Clustering', 'ML finds hidden patterns across incidents', 'Medium', 'High', 'High', '🟢'),
    ('5', 'Corrective Action Drafting', 'AI suggests what worked for similar incidents', 'Medium', 'Medium', 'Low', '🟡'),
    ('6', 'Audit Prep Automation', 'Auto-gather evidence, flag gaps', 'Medium', 'High', 'Medium', '🟢'),
    ('7', 'Customer Policy Mapping', 'Link regulations to your specific SOPs', 'Medium', 'Medium', 'Low', '🟡'),
    ('8', 'Questionnaire Automation', 'AI drafts 80% of compliance questionnaires', 'Low-Med', 'Medium', 'Low', '🟡'),
    ('9', 'Smart Onboarding', 'AI-guided customer setup', 'Low', 'Medium', 'Low', '🟡'),
    ('10', 'Hazard Detection', 'CV spots hazards in photos/video', 'High', 'High', 'High', '🟢'),
]
for i, row_data in enumerate(matrix_data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

# Strategic Positioning
doc.add_paragraph()
doc.add_heading('Strategic Positioning', level=1)

doc.add_heading('Where We Can Lead (Differentiation)', level=2)
doc.add_paragraph('• Root Cause Clustering — No one is doing real ML pattern detection on incident data')
doc.add_paragraph('• Hazard Detection — Bridge between EHS software and CV; integrated workflow')
doc.add_paragraph('• Audit Prep Automation — AI-powered evidence gathering; anxiety-killer')

doc.add_heading('Where We Should Fast-Follow (Table Stakes)', level=2)
doc.add_paragraph('• Document Ingestion — Foundation everyone will need; get there first')
doc.add_paragraph('• Incident Summarization — Easy to replicate but builds momentum')
doc.add_paragraph('• Policy Interpretation — VelocityEHS has head start but shallow')

doc.add_heading('Where We Can Wait (No Urgency)', level=2)
doc.add_paragraph('• Questionnaire Automation — Dedicated vendors exist; integrate vs. build')
doc.add_paragraph('• Smart Onboarding — Nice to have, not competitive')
doc.add_paragraph('• Corrective Action Drafting — Natural extension later')

# Prioritization
doc.add_heading('Recommended Prioritization', level=1)

doc.add_heading('🏆 Tier 1: Build Now', level=2)
doc.add_paragraph('• #2 Document Ingestion — Foundation for everything')
doc.add_paragraph('• #3 Incident Summarization — Quick win, builds confidence')
doc.add_paragraph('• #10 Hazard Detection — Already in progress; competitive race')

doc.add_heading('⭐ Tier 2: Build Next', level=2)
doc.add_paragraph('• #4 Root Cause Clustering — Strategic differentiation')
doc.add_paragraph('• #1 Policy Interpretation — High value once docs are ingested')
doc.add_paragraph('• #6 Audit Prep Automation — High customer demand')

doc.add_heading('📝 Tier 3: Evaluate Later', level=2)
doc.add_paragraph('• #5 Corrective Action Drafting')
doc.add_paragraph('• #7 Customer Policy Mapping')
doc.add_paragraph('• #8 Questionnaire Automation')
doc.add_paragraph('• #9 Smart Onboarding')

# Footer
doc.add_paragraph()
p = doc.add_paragraph('Novara Product Team | February 2026 | Confidential')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.italic = True

# Save
output_path = r'C:\Users\Grant\.openclaw\workspace-private\novara-ai\ai-opportunities-exec-summary.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
